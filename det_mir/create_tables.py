import docx

from docx import Document
from docx.shared import Cm, Pt


def keep_table_on_one_page(doc):
    tags = doc.element.xpath('//w:tr[position() < last()]/w:tc/w:p')
    for tag in tags:
        ppr = tag.get_or_add_pPr()
        ppr.keepNext_val = True


def add_style_to_table(table, size=22, WD_ALIGN_VERTICAL=None):  # центрирует текст во всех ячейках
    font = table.style.font
    font.name = 'Liberation Serif'
    font.size = Pt(size)
    for col in table.columns:
        for cell in col.cells:
            cell.paragraphs[0].alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
            # cell.vertical_alignment = docx.enum.table.WD_CELL_VERTICAL_ALIGNMENT


def create_high_table(document, supplier_name, order_num, destination, current_position):
    table = document.add_table(1, 1)
    table_params = (supplier_name,
                    order_num,
                    destination,
                    current_position)
    table.style = 'Table Grid'
    # Установка ширины таблицы == 16 см
    table.rows[0].cells[0].text = 'Маркетплейс'
    # Заполнение таблицы данными поставщика и тд.
    for param in table_params:
        row_cell = table.add_row().cells
        row_cell[0].text = param
        # row_cell[0].paragraphs[0].alignment = \
        #     docx.enum.text.WD_ALIGN_VERTICAL.CENTER

    table.columns[0].width = Cm(16)
    add_style_to_table(table)

    for row in table.rows:
        row.height = Cm(0.8)
    keep_table_on_one_page(document)


def create_bottom_table(document, barcode, model_number, product_name, quantity='1', break_flag=True):
    new_table = document.add_table(rows=2, cols=4)
    new_table.style = 'Table Grid'
    new_table.rows[0].cells[0].text = 'Код товара'
    new_table.rows[0].cells[1].text = 'Артикул'
    new_table.rows[0].cells[2].text = 'Наименование'
    new_table.rows[0].cells[3].text = 'Кол-во'

    # Заполнение строчки товара
    new_table.rows[1].cells[0].text = barcode
    new_table.rows[1].cells[1].text = model_number
    new_table.rows[1].cells[2].text = product_name
    new_table.rows[1].cells[3].text = quantity

    # Установка ширины ячеек
    new_table.columns[0].width = Cm(2.9)
    new_table.columns[1].width = Cm(4.3)
    new_table.columns[2].width = Cm(7.3)
    new_table.columns[3].width = Cm(1.5)

    # Установка высоты строк = 1 см
    for row in new_table.rows:
        row.height = Cm(1)
    keep_table_on_one_page(document)
    if break_flag:
        document.add_paragraph()
    add_style_to_table(new_table, size=14)


if __name__ == '__main__':
    doc = Document()
    create_bottom_table(doc, 'as', 'asdas', 'asdas', 'asdasd')
    doc.save('test.docx')
